#!/usr/bin/env python3
from __future__ import annotations
"""
SecureMeet - Complete Offline Meeting Transcription System
A comprehensive solution for real-time meeting transcription with speaker identification,
emotion detection, and automated meeting minutes generation.

Author: SecureMeet Development Team
License: MIT
Version: 1.0.0
"""

import sys
import os
import json
import time
import threading
import queue
import logging
import configparser
import argparse
import importlib
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, asdict
from enum import Enum
import hashlib
import base64
import struct
import sqlite3
import shutil
import subprocess
import tempfile
import uuid

def optional_import(module_name: str):
    try:
        return importlib.import_module(module_name)
    except ImportError:
        return None


pyaudio = optional_import("pyaudio")
np = optional_import("numpy")
sf = optional_import("soundfile")
signal = optional_import("scipy.signal")
whisper = optional_import("whisper")
torch = optional_import("torch")
torchaudio = optional_import("torchaudio")
pyannote_audio = optional_import("pyannote.audio")
Pipeline = getattr(pyannote_audio, "Pipeline", None) if pyannote_audio else None
Annotation = None  # reserved for future use
rnnoise = optional_import("rnnoise")

tf = optional_import("tensorflow")
sklearn_preproc = optional_import("sklearn.preprocessing")
StandardScaler = getattr(sklearn_preproc, "StandardScaler", None) if sklearn_preproc else None
joblib = optional_import("joblib")
librosa = optional_import("librosa")
_transformers = optional_import("transformers")
pipeline = getattr(_transformers, "pipeline", None) if _transformers else None

docx = optional_import("docx")
WD_ALIGN_PARAGRAPH = None
if docx is not None:
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception:
        WD_ALIGN_PARAGRAPH = None
markdown = optional_import("markdown")

cryptography_fernet = optional_import("cryptography.fernet")
Fernet = getattr(cryptography_fernet, "Fernet", None) if cryptography_fernet else None

tk = optional_import("tkinter")
ttk = messagebox = filedialog = simpledialog = None
if tk is not None:
    try:
        from tkinter import ttk, messagebox, filedialog, simpledialog
    except Exception:
        pass
plt = optional_import("matplotlib.pyplot")
FigureCanvasTkAgg = Figure = animation = None
if plt is not None:
    try:
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
        from matplotlib.figure import Figure
        import matplotlib.animation as animation
    except Exception:
        pass

DEPENDENCY_GROUPS = {
    "core": {"numpy": np, "cryptography": Fernet},
    "audio": {"pyaudio": pyaudio, "soundfile": sf, "scipy": signal, "librosa": librosa},
    "ai": {
        "tensorflow": tf,
        "scikit-learn": StandardScaler,
        "joblib": joblib,
        "openai-whisper": whisper,
        "torch": torch,
        "torchaudio": torchaudio,
        "pyannote.audio": Pipeline,
    },
    "docs": {"python-docx": docx, "markdown": markdown},
    "gui": {"tkinter": tk, "matplotlib": plt},
}


def missing_dependencies(groups: List[str]) -> List[str]:
    missing = []
    for group in groups:
        for package_name, module_obj in DEPENDENCY_GROUPS[group].items():
            if module_obj is None:
                missing.append(package_name)
    return sorted(set(missing))

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('securemeet.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Configuration constants
SAMPLE_RATE = 16000
CHUNK_SIZE = 1024
AUDIO_FORMAT = pyaudio.paInt16 if pyaudio else None
CHANNELS = 1
BUFFER_DURATION = 2.0  # seconds
OVERLAP_DURATION = 0.5  # seconds
MAX_SPEAKERS = 15
EMOTION_CLASSES = ['neutral', 'calm', 'happy', 'sad', 'angry', 'fearful', 'disgust', 'surprised']
SUPPORTED_LANGUAGES = ['en', 'es', 'fr', 'de', 'it', 'pt', 'ru', 'ja', 'ko', 'zh']

@dataclass
class AudioSegment:
    data: np.ndarray
    timestamp: float
    duration: float
    sample_rate: int
    speaker_id: Optional[str] = None
    emotion: Optional[str] = None
    confidence: Optional[float] = None

@dataclass
class TranscriptSegment:
    text: str
    start_time: float
    end_time: float
    speaker_id: str
    emotion: str
    emotion_confidence: float
    language: str
    confidence: float

@dataclass
class MeetingMinutes:
    title: str
    date: datetime
    attendees: List[str]
    duration: float
    transcript: List[TranscriptSegment]
    action_items: List[Dict[str, Any]]
    decisions: List[Dict[str, Any]]
    risks: List[Dict[str, Any]]
    summary: str

# ----------------------------------------------------------------------------------
# EmotionDetector
# ----------------------------------------------------------------------------------
class EmotionDetector:
    def __init__(self, model_path: Optional[str] = None):
        self.model = None
        self.scaler = None
        self.is_loaded = False
        self.emotion_classes = EMOTION_CLASSES
        if model_path and os.path.exists(model_path):
            self.load_model(model_path)
        else:
            self.create_model()

    def create_model(self):
        if tf is None or StandardScaler is None:
            logger.warning("Emotion detection disabled: tensorflow/scikit-learn missing.")
            return
        try:
            self.model = tf.keras.Sequential([
                tf.keras.layers.Conv1D(128, 5, activation='relu', input_shape=(40, 1)),
                tf.keras.layers.MaxPooling1D(2),
                tf.keras.layers.Conv1D(128, 5, activation='relu'),
                tf.keras.layers.MaxPooling1D(2),
                tf.keras.layers.Conv1D(128, 5, activation='relu'),
                tf.keras.layers.GlobalMaxPooling1D(),
                tf.keras.layers.Dense(128, activation='relu'),
                tf.keras.layers.Dropout(0.3),
                tf.keras.layers.Dense(len(self.emotion_classes), activation='softmax')
            ])
            self.model.compile(optimizer='adam',
                               loss='sparse_categorical_crossentropy',
                               metrics=['accuracy'])
            self.scaler = StandardScaler()
            self.is_loaded = True
            logger.info("Created emotion detection model")
        except Exception as e:
            logger.error(f"Failed to create emotion model: {e}")
            self.is_loaded = False

    def extract_features(self, audio_data: np.ndarray, sample_rate: int) -> np.ndarray:
        try:
            mfccs = librosa.feature.mfcc(
                y=audio_data.astype(float),
                sr=sample_rate,
                n_mfcc=40,
                n_fft=2048,
                hop_length=512
            ).T
            target_length = 40
            if mfccs.shape[0] > target_length:
                mfccs = mfccs[:target_length]
            else:
                mfccs = np.pad(mfccs, ((0, target_length - mfccs.shape[0]), (0, 0)), mode='constant')
            return mfccs
        except Exception as e:
            logger.error(f"Feature extraction failed: {e}")
            return np.zeros((40, 40))

    def predict_emotion(self, audio_data: np.ndarray, sample_rate: int) -> Tuple[str, float]:
        if not self.is_loaded:
            return "neutral", 0.0
        try:
            feat = self.extract_features(audio_data, sample_rate).reshape(1, -1, 1)
            pred = self.model.predict(feat, verbose=0)
            idx = int(np.argmax(pred[0]))
            return self.emotion_classes[idx], float(pred[0][idx])
        except Exception as e:
            logger.error(f"Emotion prediction failed: {e}")
            return "neutral", 0.0

    def load_model(self, model_path: str):
        try:
            self.model = tf.keras.models.load_model(model_path)
            s_path = model_path.replace('.h5', '_scaler.joblib')
            if os.path.exists(s_path):
                self.scaler = joblib.load(s_path)
            self.is_loaded = True
            logger.info(f"Loaded emotion model from {model_path}")
        except Exception as e:
            logger.error(f"Failed to load emotion model: {e}")
            self.create_model()

    def save_model(self, model_path: str):
        try:
            if self.model:
                self.model.save(model_path)
                if self.scaler:
                    joblib.dump(self.scaler, model_path.replace('.h5', '_scaler.joblib'))
                logger.info(f"Saved emotion model to {model_path}")
        except Exception as e:
            logger.error(f"Failed to save emotion model: {e}")

# ----------------------------------------------------------------------------------
# AudioProcessor
# ----------------------------------------------------------------------------------
class AudioProcessor:
    def __init__(self):
        self.sample_rate = SAMPLE_RATE
        self.chunk_size = CHUNK_SIZE
        self.audio_format = AUDIO_FORMAT
        self.channels = CHANNELS
        self.is_recording = False
        self.audio_queue = queue.Queue()
        self.noise_suppressor = None
        self.pyaudio_instance = None
        self.stream = None
        if rnnoise is not None:
            try:
                self.noise_suppressor = rnnoise.RNNoise()
                logger.info("RNNoise initialized")
            except Exception as e:
                logger.warning(f"RNNoise init failed: {e}")
        else:
            logger.info("RNNoise not available; skipping noise suppression")

    def initialize_audio(self) -> bool:
        if pyaudio is None or np is None:
            logger.error("Audio runtime dependencies are missing (pyaudio/numpy).")
            return False
        try:
            self.pyaudio_instance = pyaudio.PyAudio()
            idx = self.find_best_input_device()
            if idx is None:
                logger.error("No input device")
                return False
            self.stream = self.pyaudio_instance.open(
                format=self.audio_format,
                channels=self.channels,
                rate=self.sample_rate,
                input=True,
                input_device_index=idx,
                frames_per_buffer=self.chunk_size,
                stream_callback=self.audio_callback
            )
            logger.info(f"Audio initialized, device {idx}")
            return True
        except Exception as e:
            logger.error(f"Audio init failed: {e}")
            return False

    def find_best_input_device(self) -> Optional[int]:
        try:
            best, best_ch = None, 0
            for i in range(self.pyaudio_instance.get_device_count()):
                info = self.pyaudio_instance.get_device_info_by_index(i)
                if info['maxInputChannels'] > best_ch:
                    best, best_ch = i, info['maxInputChannels']
            return best
        except Exception as e:
            logger.error(f"find device failed: {e}")
            return None

    def audio_callback(self, in_data, frame_count, time_info, status):
        try:
            data = np.frombuffer(in_data, dtype=np.int16)
            if self.noise_suppressor is not None:
                float_data = data.astype(np.float32) / 32768.0
                den = self.noise_suppressor.process(float_data)
                data = (den * 32768.0).astype(np.int16)
            if not self.audio_queue.full():
                self.audio_queue.put(data)
            return (None, pyaudio.paContinue)
        except Exception as e:
            logger.error(f"Callback error: {e}")
            return (None, pyaudio.paAbort)

    def start_recording(self) -> bool:
        if not self.stream and not self.initialize_audio():
            return False
        self.stream.start_stream()
        self.is_recording = True
        logger.info("Recording started")
        return True

    def stop_recording(self):
        self.is_recording = False
        if self.stream:
            self.stream.stop_stream()
        logger.info("Recording stopped")

    def get_audio_data(self) -> Optional[np.ndarray]:
        try:
            return self.audio_queue.get_nowait() if not self.audio_queue.empty() else None
        except queue.Empty:
            return None

    def cleanup(self):
        try:
            if self.stream:
                self.stream.close()
            if self.pyaudio_instance:
                self.pyaudio_instance.terminate()
            logger.info("Audio cleaned up")
        except Exception as e:
            logger.error(f"Cleanup error: {e}")

# ----------------------------------------------------------------------------------
# SpeakerDiarization
# ----------------------------------------------------------------------------------
class SpeakerDiarization:
    def __init__(self, model_path: Optional[str] = None):
        self.pipeline = None
        self.is_loaded = False
        self.speaker_counter = 0
        self.initialize_pipeline(model_path)

    def initialize_pipeline(self, model_path: Optional[str]):
        if Pipeline is None:
            logger.warning("Speaker diarization disabled: pyannote.audio missing.")
            return
        try:
            if model_path and os.path.exists(model_path):
                self.pipeline = Pipeline.from_pretrained(model_path)
            else:
                self.pipeline = Pipeline.from_pretrained(
                    "pyannote/speaker-diarization-3.1", use_auth_token=None)
            self.is_loaded = True
            logger.info("Diarization pipeline ready")
        except Exception as e:
            logger.error(f"Diarization init failed: {e}")

    def process_audio(self, audio_data: np.ndarray, sample_rate: int):
        if not self.is_loaded:
            return {"speakers": [], "segments": []}
        try:
            tmp = tempfile.NamedTemporaryFile(suffix=".wav", delete=False)
            sf.write(tmp.name, audio_data, sample_rate)
            diar = self.pipeline(tmp.name)
            os.unlink(tmp.name)
            segs, spks = [], set()
            for turn, _, sp in diar.itertracks(yield_label=True):
                segs.append({"start": turn.start, "end": turn.end, "speaker": sp})
                spks.add(sp)
            return {"speakers": list(spks), "segments": segs}
        except Exception as e:
            logger.error(f"Diarization failed: {e}")
            return {"speakers": [], "segments": []}

# ----------------------------------------------------------------------------------
# WhisperTranscriber
# ----------------------------------------------------------------------------------
class WhisperTranscriber:
    def __init__(self, model_size="base"):
        self.model_size = model_size
        self.current_language = None
        self.language_confidence = 0.0
        self.model = None
        self.is_loaded = False
        self.load_model()

    def load_model(self):
        if whisper is None:
            logger.warning("Transcription disabled: openai-whisper missing.")
            return
        try:
            self.model = whisper.load_model(self.model_size)
            self.is_loaded = True
            logger.info(f"Whisper {self.model_size} loaded")
        except Exception as e:
            logger.error(f"Whisper load error: {e}")

    def transcribe_audio(self, audio_data: np.ndarray, sample_rate: int):
        if not self.is_loaded:
            return {"text": "", "language": "en", "confidence": 0.0}
        try:
            if audio_data.dtype != np.float32:
                audio_data = audio_data.astype(np.float32) / 32768.0
            if sample_rate != 16000:
                audio_data = librosa.resample(audio_data, sample_rate, 16000)
            res = self.model.transcribe(audio_data, language=self.current_language,
                                        fp16=False, verbose=False)
            lang = res.get("language", "en")
            if self.current_language is None or self.language_confidence < 0.9:
                self.current_language = lang
                self.language_confidence = min(1.0, self.language_confidence + 0.1)
            return {"text": res["text"].strip(), "language": lang,
                    "confidence": 1.0, "segments": res.get("segments", [])}
        except Exception as e:
            logger.error(f"Transcription error: {e}")
            return {"text": "", "language": "en", "confidence": 0.0}

# ----------------------------------------------------------------------------------
# NotesExtractor
# ----------------------------------------------------------------------------------
class NotesExtractor:
    def __init__(self):
        self.action_keywords = ["will", "shall", "must", "should", "need to",
                                "action", "task", "todo", "follow up", "next step"]
        self.decision_keywords = ["decided", "agreed", "concluded", "determined",
                                  "resolved", "decision", "consensus", "approve"]
        self.risk_keywords = ["risk", "issue", "problem", "concern", "challenge",
                              "blocker", "obstacle", "threat", "danger", "warning"]
        self.parking_keywords = ["park", "parking", "table", "defer", "postpone",
                                 "later", "offline", "separate"]

    def extract_notes(self, transcript: List[TranscriptSegment]):
        notes = {"action_items": [], "decisions": [], "risks": [], "parking_lot": []}
        for s in transcript:
            text_l = s.text.lower()
            if any(k in text_l for k in self.action_keywords):
                notes["action_items"].append({"text": s.text, "speaker": s.speaker_id,
                                              "timestamp": s.start_time, "confidence": s.confidence})
            if any(k in text_l for k in self.decision_keywords):
                notes["decisions"].append({"text": s.text, "speaker": s.speaker_id,
                                           "timestamp": s.start_time, "confidence": s.confidence})
            if any(k in text_l for k in self.risk_keywords):
                notes["risks"].append({"text": s.text, "speaker": s.speaker_id,
                                       "timestamp": s.start_time, "confidence": s.confidence})
            if any(k in text_l for k in self.parking_keywords):
                notes["parking_lot"].append({"text": s.text, "speaker": s.speaker_id,
                                             "timestamp": s.start_time, "confidence": s.confidence})
        return notes

# ----------------------------------------------------------------------------------
# DocumentExporter
# ----------------------------------------------------------------------------------
class DocumentExporter:
    def export_to_word(self, minutes: MeetingMinutes, path: str):
        if docx is None or WD_ALIGN_PARAGRAPH is None:
            logger.error("Word export unavailable: python-docx missing.")
            return
        try:
            doc = docx.Document()
            h = doc.add_heading(minutes.title, 0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_heading('Meeting Information', 1)
            tbl = doc.add_table(rows=4, cols=2)
            tbl.style = 'Table Grid'
            tbl.cell(0, 0).text = 'Date';        tbl.cell(0, 1).text = minutes.date.strftime('%Y-%m-%d %H:%M')
            tbl.cell(1, 0).text = 'Duration';    tbl.cell(1, 1).text = f"{minutes.duration:.1f} minutes"
            tbl.cell(2, 0).text = 'Attendees';   tbl.cell(2, 1).text = ', '.join(minutes.attendees)
            tbl.cell(3, 0).text = 'Summary';     tbl.cell(3, 1).text = minutes.summary
            if minutes.action_items:
                doc.add_heading('Action Items', 1)
                for i, it in enumerate(minutes.action_items, 1):
                    doc.add_paragraph(f"{i}. {it['text']} (Speaker: {it['speaker']})")
            if minutes.decisions:
                doc.add_heading('Decisions', 1)
                for i, d in enumerate(minutes.decisions, 1):
                    doc.add_paragraph(f"{i}. {d['text']} (Speaker: {d['speaker']})")
            if minutes.risks:
                doc.add_heading('Risks & Issues', 1)
                for i, r in enumerate(minutes.risks, 1):
                    doc.add_paragraph(f"{i}. {r['text']} (Speaker: {r['speaker']})")
            doc.add_heading('Full Transcript', 1)
            for s in minutes.transcript:
                p = doc.add_paragraph()
                p.add_run(f"[{s.start_time:.1f}s] {s.speaker_id}: ").bold = True
                p.add_run(s.text)
                p.add_run(f" ({s.emotion})").italic = True
            doc.save(path)
            logger.info(f"Word saved {path}")
        except Exception as e:
            logger.error(f"Word export error: {e}")

    def export_to_markdown(self, minutes: MeetingMinutes, path: str):
        try:
            md = [f"# {minutes.title}",
                  "## Meeting Information",
                  f"- **Date**: {minutes.date.strftime('%Y-%m-%d %H:%M')}",
                  f"- **Duration**: {minutes.duration:.1f} minutes",
                  f"- **Attendees**: {', '.join(minutes.attendees)}",
                  "## Summary",
                  minutes.summary, ""]
            if minutes.action_items:
                md += ["## Action Items"] + [f"{i}. {it['text']} (Speaker: {it['speaker']})"
                                             for i, it in enumerate(minutes.action_items, 1)] + [""]
            if minutes.decisions:
                md += ["## Decisions"] + [f"{i}. {d['text']} (Speaker: {d['speaker']})"
                                          for i, d in enumerate(minutes.decisions, 1)] + [""]
            if minutes.risks:
                md += ["## Risks & Issues"] + [f"{i}. {r['text']} (Speaker: {r['speaker']})"
                                               for i, r in enumerate(minutes.risks, 1)] + [""]
            md += ["## Full Transcript"] + [f"**[{s.start_time:.1f}s] {s.speaker_id}**: {s.text} *({s.emotion})*"
                                             for s in minutes.transcript]
            Path(path).write_text("\n".join(md), encoding="utf-8")
            logger.info(f"Markdown saved {path}")
        except Exception as e:
            logger.error(f"Markdown export error: {e}")

    def export_to_text(self, minutes: MeetingMinutes, path: str):
        try:
            txt = [minutes.title, "=" * len(minutes.title),
                   "Meeting Information:",
                   f"Date: {minutes.date.strftime('%Y-%m-%d %H:%M')}",
                   f"Duration: {minutes.duration:.1f} minutes",
                   f"Attendees: {', '.join(minutes.attendees)}", "",
                   "Summary:", minutes.summary, ""]
            if minutes.action_items:
                txt += ["Action Items:"] + [f"{i}. {it['text']} (Speaker: {it['speaker']})"
                                            for i, it in enumerate(minutes.action_items, 1)] + [""]
            if minutes.decisions:
                txt += ["Decisions:"] + [f"{i}. {d['text']} (Speaker: {d['speaker']})"
                                         for i, d in enumerate(minutes.decisions, 1)] + [""]
            if minutes.risks:
                txt += ["Risks & Issues:"] + [f"{i}. {r['text']} (Speaker: {r['speaker']})"
                                              for i, r in enumerate(minutes.risks, 1)] + [""]
            txt += ["Full Transcript:"] + [f"[{s.start_time:.1f}s] {s.speaker_id}: {s.text} ({s.emotion})"
                                           for s in minutes.transcript]
            Path(path).write_text("\n".join(txt), encoding="utf-8")
            logger.info(f"Text saved {path}")
        except Exception as e:
            logger.error(f"Text export error: {e}")

# ----------------------------------------------------------------------------------
# DataManager
# ----------------------------------------------------------------------------------
class DataManager:
    def __init__(self, data_dir="data"):
        if Fernet is None:
            raise RuntimeError("cryptography is required for secure storage.")
        self.data_dir = Path(data_dir)
        self.data_dir.mkdir(exist_ok=True)
        self.encryption_key = self.get_or_create_key()
        self.fernet = Fernet(self.encryption_key)
        self.db_path = self.data_dir / "meetings.db"
        self.init_database()

    def get_or_create_key(self):
        kf = self.data_dir / "encryption.key"
        if kf.exists():
            return kf.read_bytes()
        key = Fernet.generate_key()
        kf.write_bytes(key)
        return key

    def init_database(self):
        try:
            with sqlite3.connect(self.db_path) as c:
                c.execute('''CREATE TABLE IF NOT EXISTS meetings (
                    id TEXT PRIMARY KEY,
                    title TEXT, date TEXT, duration REAL,
                    attendees TEXT, transcript_path TEXT,
                    audio_path TEXT, created_at TEXT)''')
                c.execute('''CREATE TABLE IF NOT EXISTS segments (
                    id TEXT PRIMARY KEY, meeting_id TEXT,
                    text TEXT, start_time REAL, end_time REAL,
                    speaker_id TEXT, emotion TEXT, confidence REAL,
                    FOREIGN KEY(meeting_id) REFERENCES meetings(id))''')
                c.commit()
            logger.info("DB ready")
        except Exception as e:
            logger.error(f"DB init error: {e}")

    def save_meeting(self, minutes: MeetingMinutes, audio: Optional[np.ndarray] = None):
        try:
            mid = str(uuid.uuid4())
            mdir = self.data_dir / mid
            mdir.mkdir(exist_ok=True)
            tr_path = mdir / "transcript.enc"
            tr_data = self.fernet.encrypt(json.dumps([asdict(s) for s in minutes.transcript]).encode())
            tr_path.write_bytes(tr_data)
            aud_path = None
            if audio is not None:
                aud_path = mdir / "audio.enc"
                aud_path.write_bytes(self.fernet.encrypt(audio.tobytes()))
            with sqlite3.connect(self.db_path) as c:
                c.execute('''INSERT INTO meetings VALUES (?,?,?,?,?,?,?,?)''',
                          (mid, minutes.title, minutes.date.isoformat(), minutes.duration,
                           json.dumps(minutes.attendees), str(tr_path),
                           str(aud_path) if aud_path else None, datetime.now().isoformat()))
                for s in minutes.transcript:
                    c.execute('''INSERT INTO segments VALUES (?,?,?,?,?,?,?,?)''',
                              (str(uuid.uuid4()), mid, s.text, s.start_time, s.end_time,
                               s.speaker_id, s.emotion, s.confidence))
                c.commit()
            logger.info(f"Meeting saved {mid}")
            return mid
        except Exception as e:
            logger.error(f"Save meeting error: {e}")
            return ""

    def load_meeting(self, mid: str) -> Optional[MeetingMinutes]:
        try:
            with sqlite3.connect(self.db_path) as c:
                row = c.execute('SELECT * FROM meetings WHERE id=?', (mid,)).fetchone()
                if not row:
                    return None
                tr_path = row[5]
                tr_data = json.loads(self.fernet.decrypt(Path(tr_path).read_bytes()).decode())
                transcript = [TranscriptSegment(**d) for d in tr_data]
                return MeetingMinutes(
                    title=row[1],
                    date=datetime.fromisoformat(row[2]),
                    attendees=json.loads(row[4]),
                    duration=row[3],
                    transcript=transcript,
                    action_items=[], decisions=[], risks=[], summary="")
        except Exception as e:
            logger.error(f"Load meeting error: {e}")
            return None

    def apply_retention_policy(self):
        try:
            now = datetime.now()
            aud_cut = now - timedelta(days=30)
            tr_cut = now - timedelta(days=180)
            with sqlite3.connect(self.db_path) as c:
                to_purge = c.execute('SELECT id, created_at, audio_path, transcript_path FROM meetings').fetchall()
                for mid, created, aud, tr in to_purge:
                    cdt = datetime.fromisoformat(created)
                    if aud and cdt < aud_cut:
                        try:
                            os.remove(aud)
                            c.execute('UPDATE meetings SET audio_path=NULL WHERE id=?', (mid,))
                        except Exception:
                            pass
                    if cdt < tr_cut:
                        try:
                            if tr: os.remove(tr)
                            Path(tr).parent.rmdir()
                            c.execute('DELETE FROM segments WHERE meeting_id=?', (mid,))
                            c.execute('DELETE FROM meetings WHERE id=?', (mid,))
                        except Exception:
                            pass
                c.commit()
        except Exception as e:
            logger.error(f"Retention error: {e}")

# ----------------------------------------------------------------------------------
# SecureMeetGUI
# ----------------------------------------------------------------------------------
class SecureMeetGUI:
    def __init__(self):
        missing_gui = missing_dependencies(["core", "audio", "ai", "docs", "gui"])
        if missing_gui:
            raise RuntimeError(
                "Cannot start GUI. Missing dependencies: "
                + ", ".join(missing_gui)
                + ". Run: ./install_deps.sh"
            )

        self.root = tk.Tk()
        self.root.title("SecureMeet - Offline Meeting Transcription")
        self.root.geometry("1200x800")

        self.audio_processor = AudioProcessor()
        self.transcriber = WhisperTranscriber()
        self.speaker_diarizer = SpeakerDiarization()
        self.emotion_detector = EmotionDetector()
        self.notes_extractor = NotesExtractor()
        self.document_exporter = DocumentExporter()
        self.data_manager = DataManager()

        self.is_recording = False
        self.current_transcript: List[TranscriptSegment] = []
        self.start_time = 0
        self.audio_buffer: List[np.ndarray] = []

        self.processing_thread = None
        self.stop_processing = threading.Event()

        self.setup_ui()
        self.setup_keybindings()
        self.check_retention_policy()

    # UI -------------------------------------------------------------------------
    def setup_ui(self):
        menubar = tk.Menu(self.root); self.root.config(menu=menubar)
        file_menu = tk.Menu(menubar, tearoff=0); menubar.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="New Meeting", command=self.start_new_meeting)
        file_menu.add_command(label="Open Meeting", command=self.open_meeting)
        file_menu.add_command(label="Export...", command=self.export_meeting)
        file_menu.add_separator(); file_menu.add_command(label="Exit", command=self.root.quit)

        toolbar = ttk.Frame(self.root); toolbar.pack(side=tk.TOP, fill=tk.X, padx=5, pady=5)
        self.record_button = ttk.Button(toolbar, text="Start Recording", command=self.toggle_recording)
        self.record_button.pack(side=tk.LEFT, padx=5)
        self.stop_button = ttk.Button(toolbar, text="Stop", command=self.stop_recording, state=tk.DISABLED)
        self.stop_button.pack(side=tk.LEFT, padx=5)
        self.bookmark_button = ttk.Button(toolbar, text="Bookmark", command=self.add_bookmark)
        self.bookmark_button.pack(side=tk.LEFT, padx=5)

        self.status_var = tk.StringVar(value="Ready")
        ttk.Label(self.root, textvariable=self.status_var, relief=tk.SUNKEN)\
            .pack(side=tk.BOTTOM, fill=tk.X)

        main = ttk.Frame(self.root); main.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        left = ttk.Frame(main); left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        ttk.Label(left, text="Live Transcript", font=('Arial', 12, 'bold')).pack()
        self.transcript_text = tk.Text(left, height=30, width=60, wrap=tk.WORD,
                                       font=('Arial', 10), state=tk.DISABLED)
        scr = ttk.Scrollbar(left, orient=tk.VERTICAL, command=self.transcript_text.yview)
        self.transcript_text.config(yscrollcommand=scr.set)
        self.transcript_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scr.pack(side=tk.RIGHT, fill=tk.Y)

        right = ttk.Frame(main, width=350); right.pack(side=tk.RIGHT, fill=tk.Y, padx=5)
        ttk.Label(right, text="Notes Extracted", font=('Arial', 12, 'bold')).pack()
        self.notes_tree = ttk.Treeview(right, columns=("type","speaker","text"),
                                       show="headings", height=25)
        self.notes_tree.heading("type", text="Type")
        self.notes_tree.heading("speaker", text="Speaker")
        self.notes_tree.heading("text", text="Text")
        self.notes_tree.column("type", width=80, anchor=tk.CENTER)
        self.notes_tree.column("speaker", width=60, anchor=tk.CENTER)
        self.notes_tree.column("text", width=200)
        self.notes_tree.pack(fill=tk.BOTH, expand=True)

    def setup_keybindings(self):
        self.root.bind("<F5>", lambda e: self.toggle_recording())
        self.root.bind("<F6>", lambda e: self.stop_recording())
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

    # Meeting lifecycle -----------------------------------------------------------
    def start_new_meeting(self):
        if self.is_recording:
            messagebox.showwarning("Recording", "Stop current recording first.")
            return
        title = simpledialog.askstring("New Meeting", "Enter meeting title:")
        if not title: return
        attendees = simpledialog.askstring("Attendees", "Comma-separated attendee list:")
        self.meeting_meta = {"title": title.strip(),
                             "attendees": [a.strip() for a in attendees.split(",")] if attendees else []}
        self.current_transcript.clear(); self.audio_buffer.clear()
        self.transcript_text.config(state=tk.NORMAL); self.transcript_text.delete("1.0", tk.END)
        self.transcript_text.config(state=tk.DISABLED)
        for row in self.notes_tree.get_children(): self.notes_tree.delete(row)
        self.status_var.set(f"New meeting: {title}")
        self.start_time = time.time()
        self.toggle_recording()

    def open_meeting(self):
        mid = simpledialog.askstring("Open Meeting", "Enter meeting ID:")
        if not mid: return
        minutes = self.data_manager.load_meeting(mid)
        if not minutes:
            messagebox.showerror("Error", "Meeting not found."); return
        self.show_minutes(minutes)

    def export_meeting(self):
        if not self.current_transcript:
            messagebox.showinfo("Export", "No meeting loaded."); return
        notes = self.notes_extractor.extract_notes(self.current_transcript)
        dur = self.current_transcript[-1].end_time if self.current_transcript else 0
        minutes = MeetingMinutes(
            title=self.meeting_meta["title"],
            date=datetime.fromtimestamp(self.start_time),
            attendees=self.meeting_meta["attendees"],
            duration=dur/60,
            transcript=self.current_transcript,
            action_items=notes["action_items"],
            decisions=notes["decisions"],
            risks=notes["risks"],
            summary=simpledialog.askstring("Summary", "Write short summary:") or ""
        )
        f = filedialog.asksaveasfilename(defaultextension=".docx",
                                         filetypes=[("Word","*.docx"),("Markdown","*.md"),("Text","*.txt")])
        if not f: return
        if f.endswith(".docx"):
            self.document_exporter.export_to_word(minutes, f)
        elif f.endswith(".md"):
            self.document_exporter.export_to_markdown(minutes, f)
        else:
            self.document_exporter.export_to_text(minutes, f)
        self.data_manager.save_meeting(minutes, np.concatenate(self.audio_buffer) if self.audio_buffer else None)
        messagebox.showinfo("Export", f"Saved to {f}")

    # Recording controls ----------------------------------------------------------
    def toggle_recording(self):
        if not self.is_recording:
            if self.audio_processor.start_recording():
                self.is_recording = True
                self.record_button.config(text="Pause")
                self.stop_button.config(state=tk.NORMAL)
                self.status_var.set("Recording…")
                self.stop_processing.clear()
                self.processing_thread = threading.Thread(target=self.process_audio_loop, daemon=True)
                self.processing_thread.start()
        else:
            self.is_recording = False
            self.record_button.config(text="Resume")
            self.status_var.set("Paused")

    def stop_recording(self):
        if not self.is_recording and not self.processing_thread:
            return
        self.is_recording = False
        self.audio_processor.stop_recording()
        self.stop_processing.set()
        if self.processing_thread:
            self.processing_thread.join()
        self.record_button.config(text="Start Recording")
        self.stop_button.config(state=tk.DISABLED)
        self.status_var.set("Stopped")

    def add_bookmark(self):
        if not self.is_recording: return
        ts = time.time() - self.start_time
        self.transcript_text.config(state=tk.NORMAL)
        self.transcript_text.insert(tk.END, f"\n--- Bookmark @ {ts:.1f}s ---\n")
        self.transcript_text.config(state=tk.DISABLED)

    # Audio processing loop -------------------------------------------------------
    def process_audio_loop(self):
        while not self.stop_processing.is_set():
            data = self.audio_processor.get_audio_data()
            if data is None:
                time.sleep(0.05); continue
            self.audio_buffer.append(data.copy())
            diar = self.speaker_diarizer.process_audio(data, SAMPLE_RATE)
            for seg in diar["segments"]:
                start = seg["start"]; end = seg["end"]; spk = seg["speaker"]
                segment_audio = data[int(start*SAMPLE_RATE):int(end*SAMPLE_RATE)]
                emotion, emo_conf = self.emotion_detector.predict_emotion(segment_audio, SAMPLE_RATE)
                tx = self.transcriber.transcribe_audio(segment_audio, SAMPLE_RATE)
                if not tx["text"]: continue
                speaker_id = f"S{spk}" if isinstance(spk, int) else spk
                tseg = TranscriptSegment(
                    text=tx["text"],
                    start_time=time.time()-self.start_time,
                    end_time=time.time()-self.start_time,
                    speaker_id=speaker_id,
                    emotion=emotion,
                    emotion_confidence=emo_conf,
                    language=tx["language"],
                    confidence=tx["confidence"])
                self.current_transcript.append(tseg)
                self.update_transcript_text(tseg)
                self.update_notes_view(tseg)

    # UI updates ------------------------------------------------------------------
    def update_transcript_text(self, seg: TranscriptSegment):
        self.transcript_text.config(state=tk.NORMAL)
        self.transcript_text.insert(tk.END,
            f"[{seg.start_time:.1f}s] {seg.speaker_id}: {seg.text} ({seg.emotion})\n")
        self.transcript_text.see(tk.END)
        self.transcript_text.config(state=tk.DISABLED)

    def update_notes_view(self, seg: TranscriptSegment):
        n = self.notes_extractor.extract_notes([seg])
        for typ in ("action_items","decisions","risks","parking_lot"):
            for item in n[typ]:
                self.notes_tree.insert("", tk.END,
                                       values=(typ[:-1].title(), seg.speaker_id, seg.text))

    # Retention -------------------------------------------------------------------
    def check_retention_policy(self):
        self.data_manager.apply_retention_policy()
        self.root.after(86_400_000, self.check_retention_policy)  # 24h

    # Utils -----------------------------------------------------------------------
    def show_minutes(self, minutes: MeetingMinutes):
        self.transcript_text.config(state=tk.NORMAL); self.transcript_text.delete("1.0", tk.END)
        for s in minutes.transcript:
            self.transcript_text.insert(tk.END,
                f"[{s.start_time:.1f}s] {s.speaker_id}: {s.text} ({s.emotion})\n")
        self.transcript_text.config(state=tk.DISABLED)
        for r in self.notes_tree.get_children(): self.notes_tree.delete(r)
        for typ,label in (("action_items","Action"),("decisions","Decision"),("risks","Risk")):
            for it in getattr(minutes, typ):
                self.notes_tree.insert("", tk.END, values=(label, it["speaker"], it["text"]))
        self.status_var.set(f"Loaded meeting: {minutes.title}")

    def on_close(self):
        if self.is_recording and not messagebox.askyesno("Exit","Recording in progress. Exit anyway?"):
            return
        self.stop_recording()
        self.audio_processor.cleanup()
        self.root.destroy()

    # Main loop -------------------------------------------------------------------
    def run(self):
        self.root.mainloop()

# ----------------------------------------------------------------------------------
def main(argv: list[str] | None = None) -> None:
    """Entry point for SecureMeet."""
    argv = argv or sys.argv[1:]

    parser = argparse.ArgumentParser(description="SecureMeet offline transcription tool")
    parser.add_argument("--no-gui", action="store_true", help="Run without starting the GUI")
    parser.add_argument("--check-deps", action="store_true", help="Show dependency status and exit")
    parser.add_argument("--version", action="version", version="SecureMeet 1.0.0")
    args = parser.parse_args(argv)

    if args.check_deps:
        print("SecureMeet dependency status:")
        for group in ("core", "audio", "ai", "docs", "gui"):
            missing = missing_dependencies([group])
            status = "OK" if not missing else f"missing: {', '.join(missing)}"
            print(f"- {group:5s} -> {status}")
        return

    if args.no_gui:
        logger.info("GUI disabled via --no-gui; exiting.")
        return

    if os.environ.get("DISPLAY", "") == "":
        logger.error("No display detected. Run with --no-gui or configure a DISPLAY.")
        return

    try:
        SecureMeetGUI().run()
    except RuntimeError as exc:
        logger.error(str(exc))
        logger.info("Tip: run `python3 securemeet.py --check-deps` for details.")

if __name__ == "__main__":
    main()
